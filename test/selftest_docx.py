# -*- coding: utf-8 -*-
"""Word 读写核心自测（python-docx 构造含批注文档 → 读写往返）。

覆盖：read_docx 段落/批注/修订解析、write_docx 段落重建 + 批注保留、
无批注文档写回、修订标记检测。
"""
import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

PASS = []
FAIL = []


def check(name, cond, detail=""):
    if cond:
        PASS.append(name)
    else:
        FAIL.append(f"{name} {detail}")


from docx import Document
from docx.oxml.ns import qn
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.oxml import parse_xml
from lxml import etree

from src.core.docx_io import DocxComment, read_docx, write_docx


def make_docx_with_comments(path: str, paragraphs: list[str],
                            comments: list[tuple[str, str, str]],
                            ranges: list[tuple[int, int]] | None = None) -> None:
    """构造含批注的测试文档。

    comments: [(author, text, paragraph_index)]
    ranges: 每个批注的段内字符范围 (start, end)；None = 批注锚定到段落开头
    """
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)

    xml_parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">',
    ]
    for i, (author, text, _pi) in enumerate(comments):
        xml_parts.append(
            f'<w:comment w:id="{i}" w:author="{author}" w:date="2026-08-01T10:00:00Z">'
            f"<w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:comment>"
        )
    xml_parts.append("</w:comments>")
    comments_xml = parse_xml("".join(xml_parts).encode("utf-8"))
    partname = PackURI("/word/comments.xml")
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
    part = Part(partname, content_type, etree.tostring(comments_xml), doc.part.package)
    doc.part.package.parts.append(part)
    doc.part.relate_to(
        part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
    )
    for pid, para in enumerate(doc.paragraphs):
        start = para._p.makeelement(qn("w:commentRangeStart"), {qn("w:id"): str(pid)})
        end = para._p.makeelement(qn("w:commentRangeEnd"), {qn("w:id"): str(pid)})
        ref = para._p.makeelement(qn("w:r"), {})
        ref_el = para._p.makeelement(qn("w:commentReference"), {qn("w:id"): str(pid)})
        ref.append(ref_el)
        para._p.insert(0, start)
        para._p.append(end)
        para._p.append(ref)
    doc.save(path)


tmp = Path(tempfile.mkdtemp(prefix="paperwb_docx_test_"))
try:
    # ---------- 1. 读取含批注文档 ----------
    path = tmp / "with_comments.docx"
    make_docx_with_comments(
        str(path),
        ["第一段：鸟类羽色发育的黑色素细胞研究。",
         "第二段：单细胞测序揭示了新的调控通路。"],
        [("导师", "这里需要补充参考文献", 0),
         ("导师", "口语化，建议改写", 1)],
    )
    content = read_docx(str(path))
    check("段落数", len(content.paragraphs) == 2, repr(content.paragraphs))
    check("段落文本", content.paragraphs[0].startswith("第一段"), content.paragraphs[0])
    check("批注数", len(content.comments) == 2, repr(content.comments))
    c0 = content.comments[0]
    check("批注作者", c0.author == "导师", c0.author)
    check("批注文本", c0.text == "这里需要补充参考文献", c0.text)
    check("批注锚定段落", c0.paragraph_index == 0, str(c0.paragraph_index))
    check("批注2锚定段落", content.comments[1].paragraph_index == 1,
          str(content.comments[1].paragraph_index))
    check("无修订标记", not content.has_revisions)
    # 批注锚定在段落开头（commentRangeStart 插在段首）→ 偏移 (0, 0)
    check("批注字符偏移存在", c0.char_start >= 0 and c0.char_end >= 0,
          f"{c0.char_start},{c0.char_end}")

    # ---------- 2. 写回：内容修改 + 批注保留 ----------
    write_docx(str(path), ["修改后的第一段内容。",
                           "第二段：单细胞测序揭示了新的调控通路。"],
               comments=content.comments)
    content2 = read_docx(str(path))
    check("写回内容", content2.paragraphs[0] == "修改后的第一段内容。",
          repr(content2.paragraphs))
    check("写回批注保留", len(content2.comments) == 2, repr(content2.comments))
    check("写回批注文本", content2.comments[0].text == "这里需要补充参考文献",
          content2.comments[0].text)
    check("写回批注锚定", content2.comments[0].paragraph_index == 0,
          str(content2.comments[0].paragraph_index))
    check("写回批注范围有效",
          content2.comments[0].char_start >= 0
          and content2.comments[0].char_end > content2.comments[0].char_start,
          f"{content2.comments[0].char_start},{content2.comments[0].char_end}")

    # ---------- 3. 无批注文档写回 ----------
    path2 = tmp / "plain.docx"
    doc = Document()
    doc.add_paragraph("只有一段。")
    doc.save(str(path2))
    write_docx(str(path2), ["只有一段。"])
    c3 = read_docx(str(path2))
    check("无批注写回", c3.paragraphs == ["只有一段。"] and c3.comments == [],
          repr(c3.paragraphs))

    # ---------- 4. 修订标记检测 ----------
    path3 = tmp / "revised.docx"
    doc = Document()
    p = doc.add_paragraph("修订前文本")
    ins = p._p.makeelement(qn("w:ins"), {})
    r = p._p.makeelement(qn("w:r"), {})
    t = p._p.makeelement(qn("w:t"), {})
    t.text = "新增内容"
    r.append(t)
    ins.append(r)
    p._p.append(ins)
    doc.save(str(path3))
    c4 = read_docx(str(path3))
    check("修订标记检测", c4.has_revisions, str(c4.has_revisions))
    check("修订文本合并", "新增内容" in c4.paragraphs[0], c4.paragraphs[0])

    # ---------- 5. DocxComment.to_dict ----------
    d = DocxComment(comment_id="x", author="A", text="T", paragraph_index=2).to_dict()
    check("批注 to_dict", d["comment_id"] == "x" and d["paragraph_index"] == 2, repr(d))

    # ---------- 6. 原位写回：格式完整保留 ----------
    from docx.shared import Pt
    fmt_path = tmp / "fmt.docx"
    fdoc = Document()
    fp0 = fdoc.add_paragraph()
    fr0 = fp0.add_run("这是加粗标题")
    fr0.bold = True
    fr0.font.size = Pt(16)
    fr0.font.name = "SimSun"
    fp1 = fdoc.add_paragraph()
    fr1 = fp1.add_run("第一段正文内容")
    fr1.italic = True
    fp1.add_run().add_picture(
        os.path.join(os.path.dirname(os.path.dirname(__file__)), "PaperWB.jpg"),
        width=Pt(24),
    )
    fdoc.add_paragraph("第二段")
    fdoc.sections[0].header.paragraphs[0].text = "我的页眉"
    ftable = fdoc.add_table(rows=2, cols=2)
    ftable.cell(0, 0).text = "A1"
    ftable.cell(0, 1).text = "B1"
    ftable.cell(1, 0).text = "A2"
    ftable.cell(1, 1).text = "B2"
    fdoc.save(str(fmt_path))

    write_docx(str(fmt_path), ["这是修改后的加粗标题", "第一段正文已修改", "第二段"])
    fdoc2 = Document(str(fmt_path))
    fparas = fdoc2.paragraphs
    check("写回段落数", len(fparas) == 3, str(len(fparas)))
    _r = fparas[0].runs[0]
    check("字符格式保留(加粗/字号/字体)",
          _r.bold and _r.font.size == Pt(16) and _r.font.name == "SimSun",
          f"bold={_r.bold} size={_r.font.size} font={_r.font.name}")
    check("字符格式后文本替换", fparas[0].text == "这是修改后的加粗标题",
          fparas[0].text)
    check("斜体保留", fparas[1].runs[0].italic and fparas[1].text == "第一段正文已修改",
           repr(fparas[1].runs[0].italic))
    check("行内图片保留", len(fdoc2.inline_shapes) == 1, str(len(fdoc2.inline_shapes)))
    check("页眉保留", fdoc2.sections[0].header.paragraphs[0].text == "我的页眉",
          fdoc2.sections[0].header.paragraphs[0].text)
    check("表格保留", len(fdoc2.tables) == 1 and fdoc2.tables[0].cell(0, 0).text == "A1",
          str(len(fdoc2.tables)))

    # 段落减少：多余段落删除，页眉仍保留
    write_docx(str(fmt_path), ["只剩一段"])
    fdoc3 = Document(str(fmt_path))
    check("段落减少", len(fdoc3.paragraphs) == 1
          and fdoc3.paragraphs[0].text == "只剩一段", str(len(fdoc3.paragraphs)))
    check("删段后页眉保留",
          fdoc3.sections[0].header.paragraphs[0].text == "我的页眉",
          fdoc3.sections[0].header.paragraphs[0].text)

    # 段落增加
    write_docx(str(fmt_path), ["第一段", "第二段", "第三段"])
    fdoc4 = Document(str(fmt_path))
    check("段落增加", len(fdoc4.paragraphs) == 3, str(len(fdoc4.paragraphs)))

    # ---------- 7. 原位写回：批注随段落保留（不重锚定） ----------
    # 复用第 1 节的带批注文档（path），原位写回后批注应保留且锚定正确
    write_docx(str(path), ["第一段：修改后的鸟类研究。",
                           "第二段：单细胞测序揭示了新的调控通路。"])
    c5 = read_docx(str(path))
    check("原位写回批注保留", len(c5.comments) == 2, str(len(c5.comments)))
    check("原位写回批注锚定",
          c5.comments[0].paragraph_index == 0 and c5.comments[1].paragraph_index == 1,
          str([c.paragraph_index for c in c5.comments]))
    check("原位写回内容", c5.paragraphs[0] == "第一段：修改后的鸟类研究。",
          c5.paragraphs[0])

    # ---------- 8. 同段多条批注写回：锚点不互相吞并 ----------
    multi_path = tmp / "multi_comments.docx"
    doc = Document()
    doc.add_paragraph("第一段：鸟类羽色发育的黑色素细胞研究。")
    doc.save(str(multi_path))
    xml_parts = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">',
        '<w:comment w:id="0" w:author="导师" w:date="2026-08-01T10:00:00Z">'
        "<w:p><w:r><w:t>批注一</w:t></w:r></w:p></w:comment>",
        '<w:comment w:id="1" w:author="导师" w:date="2026-08-01T10:00:00Z">'
        "<w:p><w:r><w:t>批注二</w:t></w:r></w:p></w:comment>",
        "</w:comments>",
    ]
    comments_part = Part(
        PackURI("/word/comments.xml"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        etree.tostring(parse_xml("".join(xml_parts).encode("utf-8"))),
        doc.part.package,
    )
    doc.part.package.parts.append(comments_part)
    doc.part.relate_to(
        comments_part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
    )
    p0 = doc.paragraphs[0]._p
    s0 = p0.makeelement(qn("w:commentRangeStart"), {qn("w:id"): "0"})
    s1 = p0.makeelement(qn("w:commentRangeStart"), {qn("w:id"): "1"})
    e0 = p0.makeelement(qn("w:commentRangeEnd"), {qn("w:id"): "0"})
    e1 = p0.makeelement(qn("w:commentRangeEnd"), {qn("w:id"): "1"})
    p0.insert(1, s0)
    p0.insert(2, s1)
    p0.append(e0)
    p0.append(e1)
    doc.save(str(multi_path))

    multi = read_docx(str(multi_path))
    check("同段批注解析", len(multi.comments) == 2, str(len(multi.comments)))
    write_docx(str(multi_path), ["修改后的同一段内容。", "新增第二段"])
    multi2 = read_docx(str(multi_path))
    check("同段批注保存保留", len(multi2.comments) == 2, str(len(multi2.comments)))
    check("同段批注范围有效",
          all(c.char_start >= 0 and c.char_end > c.char_start
              for c in multi2.comments),
          str([(c.char_start, c.char_end) for c in multi2.comments]))

finally:
    import shutil
    shutil.rmtree(tmp, ignore_errors=True)

print()
for name in PASS:
    print(f"[PASS] {name}")
for name in FAIL:
    print(f"[FAIL] {name}")
print(f"\n{len(PASS)} passed, {len(FAIL)} failed")
sys.exit(1 if FAIL else 0)
