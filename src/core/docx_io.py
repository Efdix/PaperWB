"""Word (.docx) 读写核心 —— 段落级读写 + 审阅批注解析 + 修订检测。

设计取舍（用户已确认）：
- 段落级格式：读入时保留段落文本与段落样式名，写回时按首个文本 run 继承字符格式
- 批注：只读解析展示；写回时按段落重新锚定保留（锚定到段落开头）
- 修订（track changes）：只读检测并合并文本，不写回修订
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


@dataclass
class DocxComment:
    """一条审阅批注。"""

    comment_id: str = ""
    author: str = ""
    date: str = ""
    text: str = ""
    anchor_text: str = ""      # 被批注的原文片段（尽力提取）
    paragraph_index: int = -1  # 锚定段落（0 基；-1 = 未锚定）
    char_start: int = -1       # 段内字符偏移（批注起点；-1 = 未解析）
    char_end: int = -1         # 段内字符偏移（批注终点；-1 = 未解析）

    def to_dict(self) -> dict:
        return {
            "comment_id": self.comment_id, "author": self.author,
            "date": self.date, "text": self.text,
            "anchor_text": self.anchor_text,
            "paragraph_index": self.paragraph_index,
            "char_start": self.char_start,
            "char_end": self.char_end,
        }


@dataclass
class DocxContent:
    """读入的 Word 文档内容。"""

    paragraphs: list[str] = field(default_factory=list)   # 段落文本（\n 分隔）
    styles: list[str] = field(default_factory=list)       # 每段样式名（与 paragraphs 对齐）
    comments: list[DocxComment] = field(default_factory=list)
    has_revisions: bool = False                            # 存在修订（track changes）
    path: str = ""

    def to_plain_text(self) -> str:
        return "\n".join(self.paragraphs)


def _para_text(p) -> str:
    """提取段落纯文本（含修订合并：w:ins 计入、w:del 不计入）。"""
    parts: list[str] = []
    for node in p._p.iter():
        tag = node.tag
        if tag == qn("w:t"):
            parts.append(node.text or "")
        elif tag == qn("w:tab"):
            parts.append("\t")
        elif tag in (qn("w:br"), qn("w:cr")):
            parts.append("\n")
        elif tag == qn("w:delText"):
            pass  # 已删除文本不显示
    return "".join(parts)


def _para_has_revision(p) -> bool:
    """段落是否含修订标记（w:ins / w:del）。"""
    return (p._p.find(qn("w:ins")) is not None
            or p._p.find(qn("w:del")) is not None)


def _para_style_name(p) -> str:
    try:
        return p.style.name if p.style is not None else ""
    except Exception:  # noqa: BLE001
        return ""


def read_docx(path: str | Path) -> DocxContent:
    """读取 .docx：段落文本 + 样式名 + 批注 + 修订标记。"""
    path = str(path)
    doc = Document(path)
    content = DocxContent(path=path)

    for p in doc.paragraphs:
        content.paragraphs.append(_para_text(p))
        content.styles.append(_para_style_name(p))
        if _para_has_revision(p):
            content.has_revisions = True

    content.comments = _parse_comments(doc, content.paragraphs)
    return content


def _parse_comments(doc: Document, paragraphs: list[str]) -> list[DocxComment]:
    """解析审阅批注（comments.xml + 段落 commentRangeStart/End 锚定）。

    字符级偏移：按段落 XML 子节点顺序累积 w:t 文本长度，在
    commentRangeStart/End 处记录段内偏移（与 _para_text 同一遍历口径）。
    """
    comments_part = _get_comments_part(doc)
    if comments_part is None:
        return []

    comments_xml = comments_part._element
    by_id: dict[str, dict] = {}
    for c in comments_xml.findall(qn("w:comment")):
        cid = c.get(qn("w:id"), "")
        author = c.get(qn("w:author"), "")
        date = c.get(qn("w:date"), "")
        text = "".join(t.text or "" for t in c.iter(qn("w:t")))
        by_id[cid] = {"author": author, "date": date, "text": text}

    if not by_id:
        return []

    # 段落级锚定 + 字符级偏移：遍历每段 XML，累积文本长度
    spans: dict[str, tuple[int, int, int]] = {}  # cid -> (para_idx, char_start, char_end)
    for pi, p in enumerate(doc.paragraphs):
        offset = 0
        current_start: dict[str, int] = {}
        for node in p._p.iter():
            tag = node.tag
            if tag == qn("w:commentRangeStart"):
                cid = node.get(qn("w:id"), "")
                if cid:
                    current_start[cid] = offset
            elif tag == qn("w:commentRangeEnd"):
                cid = node.get(qn("w:id"), "")
                if cid and cid in current_start:
                    spans[cid] = (pi, current_start[cid], offset)
                    current_start.pop(cid)
            elif tag == qn("w:t"):
                offset += len(node.text or "")
            elif tag == qn("w:delText"):
                pass  # 与 _para_text 同口径：已删除文本不计入
        # 未闭合的批注（仅 commentRangeStart）：锚定到段尾
        for cid, start in current_start.items():
            if cid not in spans:
                spans[cid] = (pi, start, offset)

    comments: list[DocxComment] = []
    for cid, info in by_id.items():
        span = spans.get(cid)
        if span is not None:
            pi, cs, ce = span
        else:
            # 兜底：仅段落级（老文档可能只有 commentReference 无 Range）
            pi = -1
            cs = ce = -1
            for pi2, p in enumerate(doc.paragraphs):
                for node in p._p.iter():
                    if node.tag == qn("w:commentReference"):
                        if node.get(qn("w:id")) == cid:
                            pi = pi2
                            break
                if pi >= 0:
                    break
        anchor = paragraphs[pi] if 0 <= pi < len(paragraphs) else ""
        # 锚定文本：优先取被批注的精确片段，其次截取段落前 60 字符
        if 0 <= cs < ce <= len(anchor):
            anchor_text = anchor[cs:ce][:60]
        else:
            anchor_text = anchor[:60]
        comments.append(DocxComment(
            comment_id=cid,
            author=info["author"],
            date=info["date"],
            text=info["text"],
            anchor_text=anchor_text,
            paragraph_index=pi,
            char_start=cs if 0 <= cs < ce else -1,
            char_end=ce if 0 <= cs < ce else -1,
        ))
    return comments


def _get_comments_part(doc: Document):
    """获取 comments part（关系挂在 document part 上，非 package 级）。"""
    try:
        return doc.part.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        )
    except (KeyError, ValueError):
        return None


def write_docx(path: str | Path, paragraphs: list[str],
               styles: list[str] | None = None,
               comments: list[DocxComment] | None = None) -> None:
    """原位修改写回 .docx —— 保留原文档全部格式与内容。

    关键：在**原文档**上打开并保存（不是新建文档），只替换正文段落的
    文本，段落属性（pPr：样式/缩进/行距）、字符格式（第一个 run 的
    rPr：字体/字号/加粗）、页眉页脚、表格、批注锚点、书签和行内对象尽量保留。

    - 段落数变化：多的删除，少的按默认样式追加
    - 批注锚点（commentRangeStart/End）随段落 XML 原位保留，无需重锚定
    - styles/comments 参数仅为兼容旧调用签名，不再使用
    """
    path = str(path)
    doc = Document(path)  # 原位打开：所有 part（含 comments/页眉等）保留

    old_paras = doc.paragraphs  # body 直接子级段落（不含表格内段落）
    n = len(paragraphs)
    for i, p in enumerate(old_paras):
        if i < n:
            _set_para_text(p, paragraphs[i])
        else:
            p._p.getparent().remove(p._p)  # 多余段落删除
    for i in range(len(old_paras), n):
        doc.add_paragraph(paragraphs[i])   # 新增段落（python-docx 插在 sectPr 前）

    doc.save(path)


# 段落内需保留的非文本元素（批注锚点/书签/校对标记等）
_KEEP_TAGS = frozenset({
    qn("w:pPr"), qn("w:commentRangeStart"), qn("w:commentRangeEnd"),
    qn("w:commentReference"), qn("w:bookmarkStart"), qn("w:bookmarkEnd"),
    qn("w:proofErr"), qn("w:permStart"), qn("w:permEnd"),
    qn("w:hyperlink"), qn("w:fldSimple"), qn("w:sdt"),
    qn("w:customXml"), qn("w:smartTag"),
})


def _set_para_text(p, text: str) -> None:
    """原位替换段落文本：保留 pPr 与第一个 run 的字符格式（rPr）。

    删除普通文本 run 和修订承载元素，保留批注/书签/超链接/行内对象等
    非普通文本元素，然后追加一个新 run 承载文本，复用原第一个 run 的
    rPr 以保持字体格式。
    """
    # 1. 抓取第一个 run 的字符格式（rPr）
    keep_rpr = None
    for child in p._p:
        if child.tag == qn("w:r"):
            rpr = child.find(qn("w:rPr"))
            if rpr is not None:
                from copy import deepcopy
                keep_rpr = deepcopy(rpr)
            break
    def _is_embedded_object(child) -> bool:
        if child.tag != qn("w:r"):
            return False
        return (
            child.find(f".//{qn('w:drawing')}") is not None
            or child.find(f".//{qn('w:object')}") is not None
        )

    range_starts = [child for child in p._p if child.tag == qn("w:commentRangeStart")]
    range_ends = [child for child in p._p if child.tag == qn("w:commentRangeEnd")]

    # 2. 删除文本 run，但保留行内图片/对象，避免保存普通文本时丢失媒体。
    for child in list(p._p):
        if child.tag not in _KEEP_TAGS and not _is_embedded_object(child):
            p._p.remove(child)
    # 3. 新 run 必须插入到批注范围结束标记之前，否则保存后批注会变成空锚点。
    if text:
        r = p._p.makeelement(qn("w:r"), {})
        if keep_rpr is not None:
            r.append(keep_rpr)
        t = p._p.makeelement(qn("w:t"), {})
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        if range_starts and range_ends:
            # 重新锚定到整段：多个批注不能因新文本只插在第一个 End 前而互相吞锚点。
            for marker in range_starts + range_ends:
                p._p.remove(marker)
            insert_index = 1 if len(p._p) > 0 and p._p[0].tag == qn("w:pPr") else 0
            for marker in range_starts:
                p._p.insert(insert_index, marker)
                insert_index += 1
            p._p.insert(insert_index, r)
            insert_index += 1
            for marker in range_ends:
                p._p.insert(insert_index, marker)
                insert_index += 1
        else:
            insert_index = len(p._p)
            for i, child in enumerate(p._p):
                if _is_embedded_object(child):
                    insert_index = i
                    break
            p._p.insert(insert_index, r)


def has_unsaved_changes(path: str | Path, current_text: str) -> bool:
    """判断磁盘文件与当前文本是否一致（用于未保存提示）。"""
    try:
        content = read_docx(path)
        return content.to_plain_text() != current_text
    except Exception:  # noqa: BLE001
        return True
